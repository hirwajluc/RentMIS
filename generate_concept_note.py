from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import datetime

doc = Document()

# ── Page margins ──────────────────────────────────────────────────────────────
for section in doc.sections:
    section.top_margin    = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin   = Cm(2.8)
    section.right_margin  = Cm(2.8)

# ── Colour palette ────────────────────────────────────────────────────────────
NAVY      = RGBColor(0x0D, 0x2B, 0x55)   # deep navy
TEAL      = RGBColor(0x00, 0x7A, 0x8A)   # teal accent
LIGHT_BG  = RGBColor(0xF0, 0xF6, 0xFA)   # pale blue-grey
WHITE     = RGBColor(0xFF, 0xFF, 0xFF)
DARK_TEXT = RGBColor(0x1A, 0x1A, 0x2E)
MID_GREY  = RGBColor(0x55, 0x65, 0x7A)
GREEN     = RGBColor(0x00, 0x7A, 0x4A)

def set_cell_bg(cell, hex_color):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement('w:shd')
    shd.set(qn('w:val'),   'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'),  hex_color)
    tcPr.append(shd)

def set_cell_borders(cell, color='D0DCE8'):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for side in ('top', 'left', 'bottom', 'right'):
        border = OxmlElement(f'w:{side}')
        border.set(qn('w:val'),   'single')
        border.set(qn('w:sz'),    '4')
        border.set(qn('w:space'), '0')
        border.set(qn('w:color'), color)
        tcBorders.append(border)
    tcPr.append(tcBorders)

def add_run(para, text, bold=False, italic=False, size=11, color=DARK_TEXT, underline=False):
    run = para.add_run(text)
    run.bold      = bold
    run.italic    = italic
    run.underline = underline
    run.font.size  = Pt(size)
    run.font.color.rgb = color
    run.font.name  = 'Calibri'
    return run

def heading(text, level=1, color=NAVY, size=None, align=WD_ALIGN_PARAGRAPH.LEFT, space_before=14, space_after=6):
    p = doc.add_paragraph()
    p.alignment = align
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after  = Pt(space_after)
    sizes = {1: 20, 2: 15, 3: 12}
    add_run(p, text, bold=True, size=size or sizes.get(level, 12), color=color)
    return p

def body(text, size=11, color=DARK_TEXT, align=WD_ALIGN_PARAGRAPH.JUSTIFY, space_after=6, italic=False):
    p = doc.add_paragraph()
    p.alignment = align
    p.paragraph_format.space_after  = Pt(space_after)
    p.paragraph_format.space_before = Pt(2)
    add_run(p, text, size=size, color=color, italic=italic)
    return p

def bullet(text, size=11):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.space_after  = Pt(3)
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.left_indent  = Cm(0.6)
    add_run(p, text, size=size, color=DARK_TEXT)
    return p

def divider(color_hex='007A8A', thickness='12'):
    p  = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after  = Pt(2)
    pPr = p._p.get_or_add_pPr()
    pb  = OxmlElement('w:pBdr')
    bot = OxmlElement('w:bottom')
    bot.set(qn('w:val'),   'single')
    bot.set(qn('w:sz'),    thickness)
    bot.set(qn('w:space'), '1')
    bot.set(qn('w:color'), color_hex)
    pb.append(bot)
    pPr.append(pb)
    return p

# ══════════════════════════════════════════════════════════════════════════════
# COVER BLOCK
# ══════════════════════════════════════════════════════════════════════════════
tbl = doc.add_table(rows=1, cols=1)
tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
tbl.style     = 'Table Grid'
cell = tbl.cell(0, 0)
set_cell_bg(cell, '0D2B55')
cell.width = Inches(6.4)

for t, s, c, sp_b, sp_a in [
    ('CONCEPT NOTE',          11, WHITE,     16, 2),
    ('RentMIS',               32, WHITE,     4,  4),
    ('Rental Management Information System', 14, RGBColor(0xA8,0xCF,0xE8), 2, 8),
    ('Transforming Property Management in Rwanda', 12, RGBColor(0xA8,0xCF,0xE8), 2, 16),
]:
    p = cell.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(sp_b)
    p.paragraph_format.space_after  = Pt(sp_a)
    add_run(p, t, bold=(s > 12), size=s, color=c)

# Teal accent bar
acc = doc.add_table(rows=1, cols=1)
acc.alignment = WD_TABLE_ALIGNMENT.CENTER
acc.style     = 'Table Grid'
acc_cell = acc.cell(0, 0)
set_cell_bg(acc_cell, '007A8A')
p = acc_cell.add_paragraph()
p.paragraph_format.space_before = Pt(5)
p.paragraph_format.space_after  = Pt(5)
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
add_run(p, f'Prepared: {datetime.date.today().strftime("%B %Y")}   |   Version 1.0   |   Confidential', size=9, color=WHITE)

doc.add_paragraph()  # spacer

# ══════════════════════════════════════════════════════════════════════════════
# 1. EXECUTIVE SUMMARY
# ══════════════════════════════════════════════════════════════════════════════
heading('1.  Executive Summary', level=1)
divider()
body(
    'RentMIS is a production-grade, cloud-ready Rental Management Information System designed '
    'specifically for the Rwandan real estate market. The platform digitises the full lifecycle '
    'of property rental — from landlord registration and unit setup to contract signing, payment '
    'collection, and government compliance — replacing fragmented paper-based and manual processes '
    'with a secure, auditable, and accessible digital workflow.'
)
body(
    'The system is built on proven enterprise technology (Java 21 / Spring Boot 3), integrates '
    'natively with Rwanda\'s NIDA identity system and the Rwanda Revenue Authority\'s EBM fiscal '
    'receipting platform, and processes payments through GLSPay. RentMIS is ready for immediate '
    'deployment on any Linux server and is designed to scale from a single landlord to a nationwide '
    'property registry.'
)

# ══════════════════════════════════════════════════════════════════════════════
# 2. PROBLEM STATEMENT
# ══════════════════════════════════════════════════════════════════════════════
heading('2.  Problem Statement', level=1)
divider()
body(
    'The Rwandan rental market — valued at hundreds of billions of Rwandan Francs annually — '
    'continues to operate largely without formal digital infrastructure. Key pain points include:'
)
problems = [
    ('Invisible tenancy agreements', 'Most contracts are verbal or paper-based, offering no legal protection for landlords or tenants and making enforcement difficult.'),
    ('Cash-based rent collection', 'Manual collection is prone to disputes, delayed payments, and no audit trail, creating financial exposure for property owners.'),
    ('No centralised visibility', 'Government bodies and financial institutions have no reliable data on occupancy rates, rental income levels, or landlord compliance.'),
    ('Fraud and identity risk', 'Landlords cannot easily verify tenant identity, leading to fraud and illegal subletting.'),
    ('Tax leakage', 'Rental income is under-reported due to the absence of automated fiscal receipting tied to rent transactions.'),
    ('Agent accountability gaps', 'Property agents operate informally, with no traceability of commissions, managed assets, or performance.'),
]
for title, desc in problems:
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.space_after  = Pt(4)
    p.paragraph_format.left_indent  = Cm(0.6)
    add_run(p, f'{title}: ', bold=True, size=11, color=NAVY)
    add_run(p, desc, size=11, color=DARK_TEXT)

# ══════════════════════════════════════════════════════════════════════════════
# 3. PROPOSED SOLUTION
# ══════════════════════════════════════════════════════════════════════════════
heading('3.  Proposed Solution', level=1)
divider()
body(
    'RentMIS addresses each pain point through an integrated, role-based digital platform that '
    'connects all actors in the rental ecosystem — government regulators, landlords, agents, '
    'and tenants — on a single secure system.'
)

heading('3.1  System Architecture', level=2, color=TEAL, size=13, space_before=8)
body(
    'The platform follows a standard three-tier architecture: a Spring Boot REST API backend, '
    'a MySQL relational database, and a responsive HTML/JS frontend served directly by the '
    'application server. All data is encrypted in transit (HTTPS), and every critical entity '
    'change is logged in an immutable audit trail. Contract records are cryptographically signed '
    'to ensure tamper-evidence.'
)

heading('3.2  User Roles & Capabilities', level=2, color=TEAL, size=13, space_before=8)

roles_tbl = doc.add_table(rows=5, cols=2)
roles_tbl.style = 'Table Grid'
roles_tbl.alignment = WD_TABLE_ALIGNMENT.LEFT
widths = [Inches(1.5), Inches(4.8)]

headers = ['Role', 'Key Capabilities']
for i, h in enumerate(headers):
    cell = roles_tbl.cell(0, i)
    set_cell_bg(cell, '0D2B55')
    set_cell_borders(cell, 'FFFFFF')
    cell.width = widths[i]
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_run(p, h, bold=True, size=10, color=WHITE)

rows_data = [
    ('Admin',    'Full system oversight · User account management · Tenant report review & decisions · Analytics dashboard · Audit log access'),
    ('Landlord', 'Property & unit registration · Floor-level area budget management · Contract creation · Payment tracking & reconciliation · Agent assignment · Tenant flagging'),
    ('Agent',    'Manage assigned properties on behalf of landlords · Commission tracking · Performance reporting'),
    ('Tenant',   'View active contracts & unit details · Download invoices · View payment history · Request repairs or raise issues'),
]
for r_idx, (role, caps) in enumerate(rows_data, start=1):
    bg = 'F0F6FA' if r_idx % 2 == 0 else 'FFFFFF'
    c0, c1 = roles_tbl.cell(r_idx, 0), roles_tbl.cell(r_idx, 1)
    for c in (c0, c1):
        set_cell_bg(c, bg)
        set_cell_borders(c)
    c0.width = widths[0]
    c1.width = widths[1]
    add_run(c0.paragraphs[0], role, bold=True, size=10, color=NAVY)
    add_run(c1.paragraphs[0], caps, size=10, color=DARK_TEXT)

doc.add_paragraph()

heading('3.3  Core Functional Modules', level=2, color=TEAL, size=13, space_before=6)

modules = [
    ('Property & Unit Management',   'Landlords register properties with detailed metadata (category, land use, number of floors) and define a per-floor area budget. The system enforces area limits at the unit level — issuing warnings at 80% occupancy and blocking over-allocation.'),
    ('Lease Contract Management',    'Digital contracts are generated with unique reference numbers and cryptographic signatures. Tenant screening flags individuals with verified misconduct reports or overdue payment history, making them visible but non-selectable for new contracts.'),
    ('Payment Processing',           'Supports manual payments (with receipt upload), single-period online checkout, and multi-period bulk payment via GLSPay. Webhook-driven status updates ensure real-time payment reconciliation across all linked payment records.'),
    ('PDF Invoice Generation',       'Professional invoices are generated on demand using iText 7, branded with property and landlord information, and downloadable by tenants.'),
    ('Tenant Reporting System',      'Landlords can file misconduct reports against tenants. Admins review, investigate, and verify reports. Verified reports are surfaced system-wide to inform future landlords.'),
    ('NIDA Identity Verification',   'The registration flow integrates with Rwanda\'s national ID API to auto-populate verified names and photos, reducing fraudulent account creation.'),
    ('EBM Fiscal Receipting',        'Payment completion triggers automatic fiscal receipt generation via Rwanda Revenue Authority\'s Inkomane platform, ensuring full tax compliance for rental transactions.'),
    ('Analytics & Dashboards',       'Role-specific dashboards present KPIs: occupancy rates, revenue trends, overdue payments, active contracts, and tenant statistics — all in real time.'),
    ('Audit Trail & Compliance',     'Every create, update, and delete operation on a critical entity is logged with the actor\'s identity, timestamp, and before/after values, providing a complete compliance record.'),
]
for mod, desc in modules:
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.space_after  = Pt(5)
    p.paragraph_format.left_indent  = Cm(0.6)
    add_run(p, f'{mod}: ', bold=True, size=11, color=NAVY)
    add_run(p, desc, size=11, color=DARK_TEXT)

# ══════════════════════════════════════════════════════════════════════════════
# 4. TARGET STAKEHOLDERS
# ══════════════════════════════════════════════════════════════════════════════
heading('4.  Target Stakeholders & Value Proposition', level=1)
divider()

stakeholders = [
    ('Government & Regulatory Bodies\n(RRA, RSSB, RGB, City of Kigali)',
     [
         'Real-time visibility into rental transactions and fiscal receipts via EBM integration',
         'Verified tenant and landlord identity through NIDA linkage',
         'Foundation for a national rental registry and urban planning data',
         'Automated tax compliance — closing the rental income reporting gap',
     ]),
    ('Property Owners & Landlords',
     [
         'Eliminate disputes with digital, signed contracts and full payment history',
         'Reduce rent arrears through automated payment tracking and overdue alerts',
         'Protect portfolio from problematic tenants via the verified reporting system',
         'Monitor all properties and agents from a single dashboard',
     ]),
    ('Real Estate Agents',
     [
         'Transparent commission tracking and payment records',
         'Professional tools to manage multiple landlords\' portfolios',
         'Digital paper trail that builds credibility with clients',
     ]),
    ('Tenants',
     [
         'Secure digital lease agreements accessible at any time',
         'Multiple payment options including online checkout',
         'Downloadable official invoices for expense claims or visa applications',
         'Transparent record of all payments and correspondence',
     ]),
    ('Banks & Financial Institutions',
     [
         'Verified rental income data to support loan applications and credit scoring',
         'Property valuation support through occupancy and payment history',
         'Potential integration point for mortgage and rental deposit products',
     ]),
    ('PropTech Investors & Accelerators',
     [
         'Working, deployed product — not a prototype',
         'Proven integration with Rwanda\'s national infrastructure (NIDA, EBM, GLSPay)',
         'Scalable architecture ready for regional expansion',
         'Clear monetisation paths: SaaS licensing, transaction fees, API access',
     ]),
]

for stakeholder, points in stakeholders:
    heading(stakeholder, level=2, color=TEAL, size=12, space_before=10, space_after=3)
    for pt in points:
        bullet(pt)

# ══════════════════════════════════════════════════════════════════════════════
# 5. TECHNOLOGY & SECURITY
# ══════════════════════════════════════════════════════════════════════════════
heading('5.  Technology & Security', level=1)
divider()

body('RentMIS is built entirely on open-source, enterprise-grade technologies with no vendor lock-in:')

tech_tbl = doc.add_table(rows=9, cols=2)
tech_tbl.style = 'Table Grid'
tech_tbl.alignment = WD_TABLE_ALIGNMENT.LEFT
widths2 = [Inches(2.4), Inches(3.9)]

for i, h in enumerate(['Component', 'Technology / Detail']):
    c = tech_tbl.cell(0, i)
    set_cell_bg(c, '0D2B55')
    set_cell_borders(c, 'FFFFFF')
    c.width = widths2[i]
    p = c.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_run(p, h, bold=True, size=10, color=WHITE)

tech_rows = [
    ('Runtime',              'Java 21 (LTS) — long-term support until 2031'),
    ('Framework',            'Spring Boot 3.2 — industry standard for enterprise Java'),
    ('Database',             'MySQL 8 — ACID-compliant, widely supported'),
    ('Authentication',       'JWT (RS256) with refresh token rotation'),
    ('API Security',         'Rate limiting (Bucket4j), CORS, CSRF protection, input validation'),
    ('Contract Integrity',   'HMAC-SHA256 cryptographic signing on all lease documents'),
    ('Infrastructure',       'Deployable on any Linux server; systemd service support'),
    ('Compliance',           'EBM (RRA) fiscal receipting · NIDA identity verification'),
]
for r_idx, (comp, detail) in enumerate(tech_rows, start=1):
    bg = 'F0F6FA' if r_idx % 2 == 0 else 'FFFFFF'
    c0, c1 = tech_tbl.cell(r_idx, 0), tech_tbl.cell(r_idx, 1)
    for c in (c0, c1):
        set_cell_bg(c, bg)
        set_cell_borders(c)
    c0.width = widths2[0]
    c1.width = widths2[1]
    add_run(c0.paragraphs[0], comp, bold=True, size=10, color=NAVY)
    add_run(c1.paragraphs[0], detail, size=10, color=DARK_TEXT)

doc.add_paragraph()

# ══════════════════════════════════════════════════════════════════════════════
# 6. DEPLOYMENT & SCALABILITY
# ══════════════════════════════════════════════════════════════════════════════
heading('6.  Deployment & Scalability', level=1)
divider()
body(
    'RentMIS is currently operational on a Linux production server and accessible via a public IP. '
    'Deployment requires only Java 21, Maven, and MySQL — no container orchestration or cloud '
    'subscription is needed for initial operation, keeping the entry cost minimal.'
)
body('The architecture supports straightforward horizontal scaling paths:')
for pt in [
    'Containerisation with Docker and Docker Compose for environment consistency',
    'Kubernetes deployment for auto-scaling under high load',
    'Read replicas for the MySQL database to handle reporting queries',
    'CDN delivery of static assets for nationwide latency reduction',
    'Multi-region deployment for Rwanda and East African Community expansion',
]:
    bullet(pt)

# ══════════════════════════════════════════════════════════════════════════════
# 7. MONETISATION MODEL
# ══════════════════════════════════════════════════════════════════════════════
heading('7.  Monetisation Model', level=1)
divider()
body('RentMIS supports multiple revenue streams that can be activated independently or in combination:')

mon_rows = [
    ('SaaS Subscription',       'Monthly / annual licensing fee per landlord account, tiered by number of properties or units managed.'),
    ('Transaction Fee',         'A small percentage or fixed fee per online payment processed through the platform.'),
    ('Government Licensing',    'A national licensing agreement with a regulatory body (e.g. RGB, City of Kigali) for use as an official landlord registry.'),
    ('API Access',              'Monetised API access for banks, insurance companies, and fintechs to query verified rental and tenancy data.'),
    ('Premium Features',        'Advanced analytics, multi-property consolidated reporting, and branded invoice templates as paid add-ons.'),
    ('Agent Marketplace',       'Verified agent listings and lead generation for landlords seeking property management services.'),
]
mon_tbl = doc.add_table(rows=len(mon_rows)+1, cols=2)
mon_tbl.style = 'Table Grid'
mon_tbl.alignment = WD_TABLE_ALIGNMENT.LEFT
mw = [Inches(2.0), Inches(4.3)]
for i, h in enumerate(['Revenue Stream', 'Description']):
    c = mon_tbl.cell(0, i)
    set_cell_bg(c, '007A8A')
    set_cell_borders(c, 'FFFFFF')
    c.width = mw[i]
    p = c.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_run(p, h, bold=True, size=10, color=WHITE)
for r_idx, (stream, desc) in enumerate(mon_rows, start=1):
    bg = 'F0F6FA' if r_idx % 2 == 0 else 'FFFFFF'
    c0, c1 = mon_tbl.cell(r_idx, 0), mon_tbl.cell(r_idx, 1)
    for c in (c0, c1):
        set_cell_bg(c, bg)
        set_cell_borders(c)
    c0.width = mw[0]
    c1.width = mw[1]
    add_run(c0.paragraphs[0], stream, bold=True, size=10, color=NAVY)
    add_run(c1.paragraphs[0], desc, size=10, color=DARK_TEXT)

doc.add_paragraph()

# ══════════════════════════════════════════════════════════════════════════════
# 8. CURRENT STATUS & ROADMAP
# ══════════════════════════════════════════════════════════════════════════════
heading('8.  Current Status & Roadmap', level=1)
divider()

heading('8.1  Current Status', level=2, color=TEAL, size=13, space_before=8)
for item in [
    'Core platform fully built and deployed on a production Linux server',
    'All four user roles (Admin, Landlord, Agent, Tenant) operational with tailored dashboards',
    'NIDA identity verification integrated and live',
    'GLSPay payment processing integrated with webhook reconciliation',
    'EBM fiscal receipting integration complete',
    'PDF invoice generation operational',
    'Tenant reporting and admin review workflow complete',
    'Per-floor area management enforced at unit creation',
    'Full audit logging active on all critical entities',
    'Source code hosted at github.com/hirwajluc/RentMIS',
]:
    bullet(item)

heading('8.2  Proposed Roadmap', level=2, color=TEAL, size=13, space_before=8)

road_tbl = doc.add_table(rows=5, cols=3)
road_tbl.style = 'Table Grid'
road_tbl.alignment = WD_TABLE_ALIGNMENT.LEFT
rw = [Inches(1.3), Inches(1.5), Inches(3.5)]
for i, h in enumerate(['Phase', 'Timeline', 'Deliverables']):
    c = road_tbl.cell(0, i)
    set_cell_bg(c, '0D2B55')
    set_cell_borders(c, 'FFFFFF')
    c.width = rw[i]
    p = c.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_run(p, h, bold=True, size=10, color=WHITE)

road_rows = [
    ('Phase 1\nFoundation',   'Q2 2026',   'Pilot with 5–10 landlords · User feedback collection · Performance hardening'),
    ('Phase 2\nGrowth',       'Q3 2026',   'Mobile-responsive progressive web app · SMS rent reminders · Multi-currency support'),
    ('Phase 3\nScale',        'Q4 2026',   'Docker/Kubernetes deployment · Multi-region support · Bank API integrations'),
    ('Phase 4\nExpansion',    '2027',       'East African Community localisation · Marketplace launch · National registry partnership'),
]
for r_idx, (phase, timeline, delivs) in enumerate(road_rows, start=1):
    bg = 'F0F6FA' if r_idx % 2 == 0 else 'FFFFFF'
    c0, c1, c2 = road_tbl.cell(r_idx, 0), road_tbl.cell(r_idx, 1), road_tbl.cell(r_idx, 2)
    for c in (c0, c1, c2):
        set_cell_bg(c, bg)
        set_cell_borders(c)
    c0.width, c1.width, c2.width = rw
    add_run(c0.paragraphs[0], phase, bold=True, size=10, color=NAVY)
    add_run(c1.paragraphs[0], timeline, size=10, color=DARK_TEXT)
    add_run(c2.paragraphs[0], delivs, size=10, color=DARK_TEXT)

doc.add_paragraph()

# ══════════════════════════════════════════════════════════════════════════════
# 9. WHY INVEST / PARTNER
# ══════════════════════════════════════════════════════════════════════════════
heading('9.  Why Partner or Invest in RentMIS', level=1)
divider()

for title, desc in [
    ('Proven, working product',         'RentMIS is not a concept or wireframe — it is a fully deployed, functioning system with real integrations. Risk of product failure is minimal.'),
    ('Rwanda-first design',             'Built from the ground up for the Rwandan regulatory environment: NIDA, EBM, GLSPay, Rwandan land use categories, Kinyarwanda language support.'),
    ('Open ecosystem',                  'Environment-variable-driven configuration means the platform can be rebranded, white-labelled, or extended without architectural changes.'),
    ('Low infrastructure cost',         'Runs on a single Linux server costing under $50/month at entry level. No expensive cloud licences or proprietary databases required.'),
    ('Significant market opportunity',  'Kigali alone has hundreds of thousands of rental units. A 1% market penetration at even a modest subscription fee represents a multi-million franc annual revenue opportunity.'),
    ('Alignment with Vision 2050',      'Contributes directly to Rwanda\'s digital economy goals: formalising the informal sector, increasing tax compliance, and building data infrastructure for smart city planning.'),
]:
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.space_after  = Pt(5)
    p.paragraph_format.left_indent  = Cm(0.6)
    add_run(p, f'{title}: ', bold=True, size=11, color=NAVY)
    add_run(p, desc, size=11, color=DARK_TEXT)

# ══════════════════════════════════════════════════════════════════════════════
# 10. CALL TO ACTION
# ══════════════════════════════════════════════════════════════════════════════
heading('10.  Call to Action', level=1)
divider()
body(
    'We invite government institutions, property industry stakeholders, technology investors, '
    'and development partners to engage with RentMIS at any of the following levels:'
)
for action in [
    'Pilot Partnership — Provide a cohort of landlords or properties for a structured pilot programme',
    'Regulatory Endorsement — Collaborate on defining standards for a national digital tenancy registry',
    'Investment — Fund the next phase of development, infrastructure, and market acquisition',
    'White-label Licensing — Deploy a customised version under your own brand for your market or region',
    'Technical Integration — Connect your financial, insurance, or identity systems via our open API',
]:
    bullet(action)

doc.add_paragraph()
body(
    'We welcome demonstrations, technical deep-dives, and partnership discussions at your convenience.',
    italic=True, color=MID_GREY
)

# ══════════════════════════════════════════════════════════════════════════════
# FOOTER / CONTACT
# ══════════════════════════════════════════════════════════════════════════════
divider(color_hex='0D2B55', thickness='16')

ft_tbl = doc.add_table(rows=1, cols=1)
ft_tbl.style = 'Table Grid'
ft_cell = ft_tbl.cell(0, 0)
set_cell_bg(ft_cell, 'F0F6FA')
set_cell_borders(ft_cell, '007A8A')
p = ft_cell.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_before = Pt(8)
p.paragraph_format.space_after  = Pt(4)
add_run(p, 'RentMIS  |  Contact: hirwajluc@gmail.com  |  GitHub: github.com/hirwajluc/RentMIS',
        size=10, color=NAVY, bold=True)
p2 = ft_cell.add_paragraph()
p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
p2.paragraph_format.space_after = Pt(8)
add_run(p2, f'© {datetime.date.today().year} RentMIS. All rights reserved. This document is confidential.',
        size=9, color=MID_GREY, italic=True)

# ── Save ──────────────────────────────────────────────────────────────────────
out_path = '/root/java-projects/RentMIS/RentMIS_Concept_Note.docx'
doc.save(out_path)
print(f'Saved: {out_path}')
